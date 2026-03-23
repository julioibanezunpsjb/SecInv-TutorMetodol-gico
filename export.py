# -*- coding: utf-8 -*-
"""
Mentor Epistemológico - Exportación a Word
Generación de documentos académicos en formato .docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from typing import Dict, Optional
import os


def export_to_word(datos: Dict, ruta_salida: str = None) -> Optional[str]:
    """
    Exporta los datos del proyecto a un documento Word.
    
    Args:
        datos: Diccionario con todos los datos del proyecto
        ruta_salida: Ruta donde guardar el documento (opcional)
        
    Returns:
        Ruta del archivo generado o None si falla
    """
    try:
        # Crear documento
        doc = Document()
        
        # Configurar estilos
        _configurar_estilos(doc)
        
        # ════════════════════════════════════════════════════════════════════════
        # PORTADA
        # ════════════════════════════════════════════════════════════════════════
        
        # Espaciado inicial
        for _ in range(4):
            doc.add_paragraph()
        
        # Título principal
        titulo = doc.add_paragraph()
        titulo_run = titulo.add_run(datos.get('titulo', 'Proyecto de Investigación'))
        titulo_run.bold = True
        titulo_run.font.size = Pt(18)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Subtítulo
        subtitulo = doc.add_paragraph()
        subtitulo_run = subtitulo.add_run("PROYECTO DE INVESTIGACIÓN")
        subtitulo_run.font.size = Pt(14)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Área
        if datos.get('area'):
            area = doc.add_paragraph()
            area_run = area.add_run(f"Área: {datos['area']}")
            area_run.font.size = Pt(12)
            area.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espaciado
        for _ in range(6):
            doc.add_paragraph()
        
        # Fecha
        fecha = doc.add_paragraph()
        fecha_run = fecha.add_run(datetime.now().strftime("%B %Y"))
        fecha_run.font.size = Pt(12)
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salto de página
        doc.add_page_break()
        
        # ════════════════════════════════════════════════════════════════════════
        # CONTENIDO
        # ════════════════════════════════════════════════════════════════════════
        
        # 1. PROBLEMA DE INVESTIGACIÓN
        doc.add_heading('1. PROBLEMA DE INVESTIGACIÓN', level=1)
        
        if datos.get('problema'):
            p = doc.add_paragraph()
            p.add_run('Planteamiento del problema:').bold = True
            doc.add_paragraph(datos['problema'])
        else:
            doc.add_paragraph('_' * 50)
        
        doc.add_paragraph()
        
        # 2. OBJETIVOS
        doc.add_heading('2. OBJETIVOS DE LA INVESTIGACIÓN', level=1)
        
        # Objetivo General
        doc.add_heading('2.1 Objetivo General', level=2)
        if datos.get('objetivo_general'):
            doc.add_paragraph(datos['objetivo_general'])
        else:
            doc.add_paragraph('_' * 50)
        
        doc.add_paragraph()
        
        # Objetivos Específicos
        doc.add_heading('2.2 Objetivos Específicos', level=2)
        if datos.get('objetivos_especificos'):
            objetivos = datos['objetivos_especificos'].split('\n')
            for obj in objetivos:
                if obj.strip():
                    p = doc.add_paragraph(obj.strip(), style='List Bullet')
        else:
            doc.add_paragraph('_' * 50)
        
        doc.add_paragraph()
        
        # 3. HIPÓTESIS
        doc.add_heading('3. HIPÓTESIS', level=1)
        
        if datos.get('sin_hipotesis'):
            p = doc.add_paragraph()
            p.add_run('Tipo de estudio sin hipótesis: ').bold = True
            doc.add_paragraph(datos.get('justificacion_sin_hip', 
                'El presente estudio es de carácter exploratorio/descriptivo, por lo cual no se formulan hipótesis.'))
        elif datos.get('hipotesis'):
            p = doc.add_paragraph()
            p.add_run('Hipótesis de investigación:').bold = True
            doc.add_paragraph(datos['hipotesis'])
            
            # Variables
            if datos.get('variable_independiente') or datos.get('variable_dependiente'):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('Variables:').bold = True
                
                if datos.get('variable_independiente'):
                    p = doc.add_paragraph()
                    p.add_run('• Variable independiente: ').bold = True
                    p.add_run(datos['variable_independiente'])
                
                if datos.get('variable_dependiente'):
                    p = doc.add_paragraph()
                    p.add_run('• Variable dependiente: ').bold = True
                    p.add_run(datos['variable_dependiente'])
        else:
            doc.add_paragraph('_' * 50)
        
        doc.add_paragraph()
        
        # 4. MARCO TEÓRICO
        doc.add_heading('4. MARCO TEÓRICO', level=1)
        
        # Antecedentes
        doc.add_heading('4.1 Antecedentes de la Investigación', level=2)
        if datos.get('antecedentes'):
            doc.add_paragraph(datos['antecedentes'])
        else:
            doc.add_paragraph('_' * 50)
        
        # Bases teóricas
        doc.add_heading('4.2 Bases Teóricas', level=2)
        if datos.get('bases_teoricas'):
            doc.add_paragraph(datos['bases_teoricas'])
        else:
            doc.add_paragraph('_' * 50)
        
        # Definición de términos
        doc.add_heading('4.3 Definición de Términos Básicos', level=2)
        if datos.get('definiciones'):
            doc.add_paragraph(datos['definiciones'])
        else:
            doc.add_paragraph('_' * 50)
        
        doc.add_paragraph()
        
        # 5. METODOLOGÍA
        doc.add_heading('5. DISEÑO METODOLÓGICO', level=1)
        
        # Tipo y diseño
        doc.add_heading('5.1 Tipo y Diseño de Investigación', level=2)
        tipo = datos.get('tipo_estudio', 'Por definir')
        diseño = datos.get('diseño', 'Por definir')
        doc.add_paragraph(f"Tipo de estudio: {tipo}")
        doc.add_paragraph(f"Diseño: {diseño}")
        
        # Población y muestra
        doc.add_heading('5.2 Población y Muestra', level=2)
        if datos.get('poblacion'):
            p = doc.add_paragraph()
            p.add_run('Población: ').bold = True
            p.add_run(datos['poblacion'])
        else:
            doc.add_paragraph('Población: Por definir')
        
        if datos.get('muestra'):
            p = doc.add_paragraph()
            p.add_run('Muestra: ').bold = True
            p.add_run(datos['muestra'])
        else:
            doc.add_paragraph('Muestra: Por definir')
        
        # Técnicas e instrumentos
        doc.add_heading('5.3 Técnicas e Instrumentos de Recolección de Datos', level=2)
        if datos.get('tecnicas'):
            p = doc.add_paragraph()
            p.add_run('Técnicas: ').bold = True
            p.add_run(', '.join(datos['tecnicas']))
        
        if datos.get('instrumentos'):
            p = doc.add_paragraph()
            p.add_run('Instrumentos: ').bold = True
            p.add_run(datos['instrumentos'])
        
        doc.add_paragraph()
        
        # 6. CRONOGRAMA (placeholder)
        doc.add_heading('6. CRONOGRAMA DE ACTIVIDADES', level=1)
        doc.add_paragraph('[El cronograma se desarrollará en etapas posteriores del proyecto]')
        
        doc.add_paragraph()
        
        # 7. REFERENCIAS (placeholder)
        doc.add_heading('7. REFERENCIAS BIBLIOGRÁFICAS', level=1)
        doc.add_paragraph('[Las referencias se agregarán según se desarrolle el marco teórico]')
        
        # Guardar documento
        if not ruta_salida:
            nombre_archivo = f"proyecto_investigacion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            ruta_salida = os.path.join('/tmp', nombre_archivo)
        
        # Asegurar que el directorio existe
        os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
        
        doc.save(ruta_salida)
        return ruta_salida
        
    except Exception as e:
        print(f"Error al exportar a Word: {str(e)}")
        return None


def _configurar_estilos(doc: Document):
    """Configura los estilos del documento."""
    
    # Estilo para títulos principales
    style = doc.styles['Heading 1']
    style.font.size = Pt(14)
    style.font.bold = True
    
    # Estilo para subtítulos
    style = doc.styles['Heading 2']
    style.font.size = Pt(12)
    style.font.bold = True
    
    # Estilo para texto normal
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.font.name = 'Times New Roman'
